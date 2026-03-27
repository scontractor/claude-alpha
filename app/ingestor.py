import io
from pathlib import Path

import chromadb
from sentence_transformers import SentenceTransformer

from app import config
from app.utils import clean_text, chunk_text, make_doc_id


# --- Parsers ---

def parse_pdf(path: Path) -> str:
    import pypdf
    reader = pypdf.PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return clean_text("\n".join(pages))


def parse_pdf_bytes(data: bytes) -> str:
    import pypdf
    reader = pypdf.PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return clean_text("\n".join(pages))


def parse_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return clean_text("\n".join(paragraphs))


def parse_docx_bytes(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return clean_text("\n".join(paragraphs))


def parse_text(path: Path) -> str:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="latin-1")
    return clean_text(text)


def parse_text_bytes(data: bytes) -> str:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1")
    return clean_text(text)


def _get_parser(path: Path):
    ext = path.suffix.lower()
    if ext == ".pdf":
        return parse_pdf
    if ext == ".docx":
        return parse_docx
    if ext in {".md", ".txt", ".markdown"}:
        return parse_text
    return None


# --- ChromaDB ---

def get_client() -> chromadb.PersistentClient:
    return chromadb.PersistentClient(path=str(config.CHROMA_DIR.resolve()))


def get_collection(client: chromadb.PersistentClient) -> chromadb.Collection:
    return client.get_or_create_collection(
        name=config.COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
    )


# --- Model ---

def load_embedding_model() -> SentenceTransformer:
    return SentenceTransformer(config.EMBED_MODEL)


# --- Ingestion ---

def _embed_and_upsert(filename: str, text: str, collection, embed_model: SentenceTransformer) -> int:
    chunks = chunk_text(text, config.CHUNK_SIZE, config.CHUNK_OVERLAP)
    if not chunks:
        return 0

    embeddings = embed_model.encode(
        chunks,
        batch_size=config.EMBED_BATCH_SIZE,
        show_progress_bar=False,
        normalize_embeddings=True,
    ).tolist()

    stem = Path(filename).stem
    ids = [f"{stem}_chunk{i}" for i in range(len(chunks))]
    ext = Path(filename).suffix.lower()
    metadatas = [
        {"source": filename, "chunk_index": i, "file_type": ext}
        for i in range(len(chunks))
    ]

    collection.upsert(ids=ids, embeddings=embeddings, documents=chunks, metadatas=metadatas)
    return len(chunks)


def ingest_file(path: Path, collection, embed_model: SentenceTransformer) -> int:
    parser = _get_parser(path)
    if parser is None:
        return 0
    text = parser(path)
    if not text:
        return 0
    return _embed_and_upsert(path.name, text, collection, embed_model)


def ingest_from_bytes(filename: str, data: bytes, collection, embed_model: SentenceTransformer) -> int:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text = parse_pdf_bytes(data)
    elif ext == ".docx":
        text = parse_docx_bytes(data)
    elif ext in {".md", ".txt", ".markdown"}:
        text = parse_text_bytes(data)
    else:
        return 0
    if not text:
        return 0
    return _embed_and_upsert(filename, text, collection, embed_model)


def ingest_all(data_dir: Path, collection, embed_model: SentenceTransformer) -> dict[str, int]:
    results = {}
    files = [
        f for f in data_dir.rglob("*")
        if f.is_file() and f.suffix.lower() in config.SUPPORTED_EXTENSIONS
    ]
    for path in files:
        count = ingest_file(path, collection, embed_model)
        results[path.name] = count
    return results


def list_sources(collection) -> list[str]:
    if collection.count() == 0:
        return []
    results = collection.get(include=["metadatas"])
    return sorted({m["source"] for m in results["metadatas"]})


def delete_file(filename: str, collection) -> int:
    results = collection.get(where={"source": filename}, include=[])
    ids = results["ids"]
    if ids:
        collection.delete(ids=ids)
    return len(ids)
